import io
from datetime import datetime


def export_pdf(script_text: str, script_num: int) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=3*cm,
            rightMargin=3*cm,
            topMargin=3*cm,
            bottomMargin=3*cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CosmicTitle',
            fontName='Helvetica-Bold',
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=6,
            alignment=TA_LEFT,
        )
        meta_style = ParagraphStyle(
            'CosmicMeta',
            fontName='Helvetica',
            fontSize=9,
            textColor=colors.HexColor('#888888'),
            spaceAfter=24,
        )
        body_style = ParagraphStyle(
            'CosmicBody',
            fontName='Courier',
            fontSize=11,
            leading=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
        )

        story = []
        story.append(Paragraph(f"SCRIPT #{script_num:03d}", title_style))
        story.append(Paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d')} · Cosmic Horror Script Engine", meta_style))
        story.append(Spacer(1, 0.3*cm))

        for para in script_text.split('\n\n'):
            para = para.strip()
            if para:
                safe = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe, body_style))
                story.append(Spacer(1, 0.2*cm))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    except ImportError:
        return script_text.encode('utf-8')


def export_docx(script_text: str, script_num: int) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(3)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)

        title = doc.add_heading(f'SCRIPT #{script_num:03d}', 0)
        title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        title.runs[0].font.size = Pt(18)

        meta = doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d')} · Cosmic Horror Script Engine")
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        for para in script_text.split('\n\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(11)
                p.paragraph_format.space_after = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except ImportError:
        return script_text.encode('utf-8')
